import streamlit as st
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from langchain.chains.question_answering import load_qa_chain
from langchain_google_genai import GoogleGenerativeAI
from pypdf import PdfReader
from docx import Document as DocxDocument
import google.generativeai as genai
from dotenv import load_dotenv
import xml.etree.ElementTree as ET
from urllib.parse import urljoin, urlparse

# Load environment variables from the .env file for local development
load_dotenv()

# --- Helper Functions for Data Extraction ---

def get_text_from_url(url):
    """Scrapes text content from a single URL."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        for script_or_style in soup(['script', 'style']):
            script_or_style.extract()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        return '\n'.join(chunk for chunk in chunks if chunk)
    except requests.RequestException as e:
        st.warning(f"Could not scrape {url}. Reason: {e}")
        return ""

def get_text_from_sitemap(base_url):
    """
    Attempts to find and scrape all URLs in a website's sitemap.
    If no sitemap is found, scrapes the base URL.
    """
    sitemap_url = urljoin(base_url, "/sitemap.xml")
    all_text = ""
    try:
        st.info(f"Looking for sitemap at: {sitemap_url}")
        response = requests.get(sitemap_url, timeout=10)
        response.raise_for_status()

        root = ET.fromstring(response.content)
        urls_to_scrape = [elem.text for elem in root.findall('.//{http://www.sitemaps.org/schemas/sitemap/0.9}loc')]
        
        st.info(f"Found {len(urls_to_scrape)} pages in sitemap. Scraping...")
        
        with st.progress(0, text="Scraping website...") as progress_bar:
            for i, url in enumerate(urls_to_scrape):
                all_text += get_text_from_url(url) + "\n\n"
                progress_bar.progress((i + 1) / len(urls_to_scrape))
        st.success("Sitemap processed successfully!")

    # *** FIX: Corrected the exception handling block to be relevant to web scraping ***
    except (requests.exceptions.RequestException, ET.ParseError) as e:
        st.warning(f"Could not find or parse sitemap.xml (Reason: {e}). Falling back to scraping just the main page.")
        all_text = get_text_from_url(base_url)

    return all_text


def get_text_from_pdf(pdf_file):
    """Extracts text from an uploaded PDF file."""
    try:
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return ""

def get_text_from_docx(docx_file):
    """Extracts text from an uploaded DOCX file."""
    try:
        doc = DocxDocument(docx_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return ""

def get_text_from_txt(txt_file):
    """Reads text from an uploaded TXT file."""
    try:
        return txt_file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
        return ""

# --- Core RAG Logic ---

def get_text_chunks(raw_text):
    """Splits raw text into manageable chunks."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(raw_text)
    return chunks

def get_vector_store(text_chunks):
    """Creates a FAISS vector store from text chunks."""
    if not text_chunks:
        st.warning("No text to process. Please provide data.")
        return None
    try:
        model_name = "all-MiniLM-L6-v2"
        embeddings = HuggingFaceEmbeddings(model_name=model_name)
        documents = [Document(page_content=chunk) for chunk in text_chunks]
        vector_store = FAISS.from_documents(documents, embedding=embeddings)
        return vector_store
    except Exception as e:
        st.error(f"Error creating vector store: {e}")
        return None

def get_conversational_chain():
    """Creates and configures the QA chain with a dynamic and more natural prompt."""
    person_name = st.session_state.get("person_name", "the user")
    prompt_template_str = f"""
    You are "{person_name}'s AI", a friendly and professional chatbot assistant that has deep knowledge about {person_name}.
    Your goal is to answer questions about {person_name} comprehensively and naturally, using the context provided as your source of truth.
    Do not mention that you are basing your answer on the provided text. Act as if you know this information innately.
    If the answer is not in the context, politely say "I'm sorry, I don't have specific information about that regarding {person_name}."
    Keep your answers helpful and human-like.
    Context:
    {{context}}
    Question:
    {{question}}
    Helpful Answer:
    """
    model = GoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.5)
    chain = load_qa_chain(llm=model, chain_type="stuff", prompt=st.session_state.prompt_template.from_template(prompt_template_str))
    return chain

def get_person_name(text_chunks):
    """Extracts the subject's full name from the document chunks using the LLM."""
    if not text_chunks: return "ResumAI"
    try:
        temp_vector_store = get_vector_store(text_chunks)
        if temp_vector_store is None: return "ResumAI" 
        name_query = "What is the full name of the person these documents are about? Respond with only the full name."
        docs = temp_vector_store.similarity_search(name_query, k=1)
        model = GoogleGenerativeAI(model="gemini-1.5-flash")
        chain = load_qa_chain(llm=model, chain_type="stuff")
        response = chain.invoke({"input_documents": docs, "question": name_query})
        name = response.get("output_text", "ResumAI").strip()
        return name
    except Exception as e:
        st.warning(f"Could not automatically determine the name. Using default. Error: {e}")
        return "ResumAI"

# --- Streamlit UI ---

def main():
    st.set_page_config(page_title="AI Chatbot", page_icon="🤖")

    if "vector_store" not in st.session_state: st.session_state.vector_store = None
    if "messages" not in st.session_state: st.session_state.messages = []
    if "person_name" not in st.session_state: st.session_state.person_name = "Your" 
    if "prompt_template" not in st.session_state:
        from langchain.prompts import PromptTemplate
        st.session_state.prompt_template = PromptTemplate

    with st.sidebar:
        st.subheader("Knowledge Base Configuration")
        st.markdown(
            """
            1.  **Provide Sources:** Upload documents and/or provide links to a personal website and GitHub profile.
            2.  **Website Note:** The app will try to find and scrape a `sitemap.xml`. If not found, it will scrape the single page.
            3.  **Process:** Click the "Process Sources" button.
            """
        )
        
        custom_url = st.text_input("Personal Website URL (e.g., https://yourblog.com)")
        github_username = st.text_input("GitHub Username (Optional)")

        uploaded_files = st.file_uploader(
            "Upload your documents (PDF, DOCX, TXT)",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            label_visibility="collapsed"
        )
        
        if st.button("Process Sources"):
            if "GOOGLE_API_KEY" not in st.secrets:
                st.error("Google API Key not found. Go to Settings > Secrets to add it.")
            elif not uploaded_files and not custom_url and not github_username:
                st.warning("Please provide at least one source (file, website, or GitHub).")
            else:
                with st.spinner("Processing sources... This might take a while depending on website size."):
                    genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

                    raw_text = ""
                    for file in uploaded_files:
                        if file.name.endswith(".pdf"): raw_text += get_text_from_pdf(file)
                        elif file.name.endswith(".docx"): raw_text += get_text_from_docx(file)
                        elif file.name.endswith(".txt"): raw_text += get_text_from_txt(file)
                    
                    if custom_url:
                        raw_text += get_text_from_sitemap(custom_url)

                    if github_username:
                        github_url = f"https://github.com/{github_username}"
                        st.info(f"Attempting to scrape GitHub...")
                        raw_text += get_text_from_url(github_url)
                    
                    text_chunks = get_text_chunks(raw_text)
                    st.session_state.vector_store = get_vector_store(text_chunks)
                    
                    if st.session_state.vector_store:
                        st.session_state.person_name = get_person_name(text_chunks)
                        st.session_state.messages = [] 
                        st.success(f"Ready to chat about {st.session_state.person_name}!")
                        st.rerun()

    st.header(f"Chat with {st.session_state.person_name}'s AI 🤖")
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("Ask a question..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)

        if st.session_state.vector_store is None:
            st.warning("Please process some sources first.")
        else:
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        docs = st.session_state.vector_store.similarity_search(prompt, k=3)
                        chain = get_conversational_chain()
                        response = chain({"input_documents": docs, "question": prompt}, return_only_outputs=True)
                        response_text = response["output_text"]
                        st.markdown(response_text)
                        st.session_state.messages.append({"role": "assistant", "content": response_text})
                    except Exception as e:
                        error_message = f"An error occurred: {e}"
                        st.error(error_message)
                        st.session_state.messages.append({"role": "assistant", "content": error_message})

if __name__ == '__main__':
    main()
